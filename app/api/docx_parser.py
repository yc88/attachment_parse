import os
import datetime
import requests
from app.schemas.const import MyErrorCode, HttpStatusCode
from app.schemas.parse import FileParseResult, RegexField, RegexFieldStatistical, ExtraField
from app.utils.app_exceptions import UnicornException
from app.api.content_regex import get_regex_val
from collections import OrderedDict
from docx import Document
from io import BytesIO


def parse_local_docx_file(file_path):
    """
    解析本地的 .doc .docx文件内容
    """
    try:
        # 获取文件内容和编码方式
        encoding = 'utf-8'
        # 获取文件名称
        file_name = os.path.basename(file_path)
        # 获取文件大小
        file_size = os.path.getsize(file_path)
        # 获取最后修改时间
        last_modified = datetime.datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M:%S')
        doc = Document(file_path)
        return get_return_result(doc, file_size, file_name, last_modified, encoding)
    except Exception as e:
        raise UnicornException(HttpStatusCode.HTTP_500_INTERNAL_SERVER_ERROR,
                               MyErrorCode.InvalidArgument.value,
                               e.__str__())


def parse_remote_docx_file(file_url):
    try:
        # 获取文件内容和编码方式
        encoding = 'utf-8'
        # 下载远程文件内容
        response = requests.get(file_url)
        if response.status_code == 200:
            # 获取文件名称
            file_name = os.path.basename(file_url)
            # 文件内容
            raw_content = response.text
            # 获取文件大小
            file_size = len(raw_content)
            doc = Document(BytesIO(response.content))
            return get_return_result(doc, file_size, file_name, last_modified="", encoding=encoding)
        else:
            raise UnicornException(HttpStatusCode.HTTP_500_INTERNAL_SERVER_ERROR,
                                   MyErrorCode.InvalidArgument.value,
                                   " 请求错误 ")
    except Exception as e:
        raise UnicornException(HttpStatusCode.HTTP_500_INTERNAL_SERVER_ERROR,
                               MyErrorCode.InvalidArgument.value,
                               e.__str__())


def get_docx_images(doc):
    images = []
    for rel in doc.part.rels:
        if "image" in doc.part.rels[rel].target_ref:
            images.append(doc.part.rels[rel].target_ref)
    return list(OrderedDict.fromkeys(val for val in images if val))


def get_return_result(doc, file_size, file_name, last_modified, encoding):
    images = get_docx_images(doc)
    title = doc.core_properties.title
    last_modified_time = doc.core_properties.modified
    if not last_modified_time:
        last_modified_time = ""
    if last_modified_time is not str:
        last_modified_time = str(last_modified_time)
    language = doc.core_properties.language
    if language is not str:
        language = str(language)
    # Extract content
    # 获取文件内容（包括表格、图片等）
    content = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            content.append(paragraph.text)
    # Extract tables from the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                content.append(cell.text)

    file_content = '\n'.join(content)
    regex_all, regex_all_statistical = get_regex_val(file_content)

    extra_field = ExtraField(
        date=last_modified_time,
        language=language,
        charset="",
        subject=title,
        images=images,
        link=[],
    )

    return FileParseResult(file_size=file_size, file_name=file_name, last_modified=last_modified,
                           file_content=file_content, file_encoding=encoding,
                           regex_all=regex_all, regex_all_statistical=regex_all_statistical, extra=extra_field)
